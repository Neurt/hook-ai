"""Document generation (Pillar 1 output).

Renders ATS CV markdown into downloadable files: .md, .docx (python-docx),
.pdf (reportlab). The visual target is a classic black-and-white serif resume
(centered name, single ' | ' contact line, bold uppercase sections with a
full-width rule, and two-column entry rows: **Org** | Location / *Title* | Dates).
DOCX/PDF share one block-level parse so both match.

Markdown conventions produced by CVTailor.to_ats:
  # Name
  Location | phone | email | links          (one contact line)
  ## SECTION
  **Category:** a, b, c                      (skills)
  **Left** | Right                           (entry head: bold left, plain right)
  *Sub* | Dates                              (entry sub: italic both)
  - bullet
"""
from __future__ import annotations

import io
import os
import re
from dataclasses import dataclass

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ENTRY_HEAD_RE = re.compile(r"^\*\*(.+?)\*\*\s*\|\s*(.*)$")
_ENTRY_SUB_RE = re.compile(r"^\*(?!\*)([^*]+?)\*\s*\|\s*(.*)$")


class DocgenError(RuntimeError):
    """A rendering dependency is missing or generation failed. Safe to catch in a
    web server (unlike a bare ImportError/SystemExit)."""


def write_markdown(path: str, content: str) -> str:
    directory = os.path.dirname(path)
    if directory:
        os.makedirs(directory, exist_ok=True)
    with open(path, "w", encoding="utf-8") as fh:
        fh.write(content)
    return path


@dataclass
class Block:
    kind: str  # h1 | h2 | h3 | bullet | para | entry_head | entry_sub
    text: str  # raw text / left column (may contain **bold**)
    right: str = ""  # right column for entry_head / entry_sub


def parse_cv_markdown(markdown: str) -> list[Block]:
    blocks: list[Block] = []
    for raw in markdown.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("### "):
            blocks.append(Block("h3", line[4:].strip()))
        elif line.startswith("## "):
            blocks.append(Block("h2", line[3:].strip()))
        elif line.startswith("# "):
            blocks.append(Block("h1", line[2:].strip()))
        elif line[:2] in ("- ", "* ") or line[:1] == "•":
            text = line[1:].strip() if line[:1] == "•" else line[2:].strip()
            blocks.append(Block("bullet", text))
        elif _ENTRY_HEAD_RE.match(line):
            m = _ENTRY_HEAD_RE.match(line)
            blocks.append(Block("entry_head", m.group(1).strip(), m.group(2).strip()))
        elif _ENTRY_SUB_RE.match(line):
            m = _ENTRY_SUB_RE.match(line)
            blocks.append(Block("entry_sub", m.group(1).strip(), m.group(2).strip()))
        else:
            blocks.append(Block("para", line))
    return blocks


def _strip_bold(text: str) -> str:
    return text.replace("**", "")


# ── DOCX ──────────────────────────────────────────────────────────────────
def _docx_runs(paragraph, text: str) -> None:
    """Add text, rendering **bold** spans as bold runs."""
    for i, part in enumerate(text.split("**")):
        if part:
            run = paragraph.add_run(part)
            run.bold = i % 2 == 1


def _docx_heading_rule(paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for k, v in (("w:val", "single"), ("w:sz", "6"), ("w:space", "2"), ("w:color", "000000")):
        bottom.set(qn(k), v)
    borders.append(bottom)
    p_pr.append(borders)


def render_docx_bytes(markdown: str) -> bytes:
    """Black-and-white serif resume as an in-memory .docx."""
    try:
        import docx  # python-docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise DocgenError("DOCX export needs python-docx (pip install python-docx)") from exc

    document = docx.Document()
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    section = document.sections[0]
    usable = section.page_width - section.left_margin - section.right_margin

    def two_col(left: str, right: str, italic: bool) -> None:
        para = document.add_paragraph()
        para.paragraph_format.space_after = Pt(1)
        para.paragraph_format.tab_stops.add_tab_stop(usable, WD_TAB_ALIGNMENT.RIGHT)
        left_run = para.add_run(left)
        left_run.bold = not italic
        left_run.italic = italic
        para.add_run("\t")
        right_run = para.add_run(right)
        right_run.italic = italic

    blocks = parse_cv_markdown(markdown)
    contact: list[str] = []
    seen_section = False

    def flush_contact() -> None:
        if contact:
            para = document.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run(" | ".join(contact)).font.size = Pt(9.5)
            contact.clear()

    for block in blocks:
        if block.kind == "h1":
            para = document.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(_strip_bold(block.text))
            run.font.size = Pt(22)
        elif block.kind == "para" and not seen_section:
            contact.append(block.text)
        elif block.kind == "h2":
            flush_contact()
            seen_section = True
            para = document.add_paragraph()
            para.paragraph_format.space_before = Pt(8)
            run = para.add_run(_strip_bold(block.text).upper())
            run.bold = True
            run.font.size = Pt(12)
            _docx_heading_rule(para)
        elif block.kind == "entry_head":
            two_col(block.text, block.right, italic=False)
        elif block.kind == "entry_sub":
            two_col(block.text, block.right, italic=True)
        elif block.kind == "h3":
            para = document.add_paragraph()
            _docx_runs(para, "**" + block.text + "**")
        elif block.kind == "bullet":
            para = document.add_paragraph(style="List Bullet")
            _docx_runs(para, block.text)
        else:  # para after a section (e.g. summary, skills)
            para = document.add_paragraph()
            _docx_runs(para, block.text)

    flush_contact()
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def write_docx_from_markdown(path: str, markdown: str) -> str:
    try:
        data = render_docx_bytes(markdown)
    except DocgenError as exc:
        raise SystemExit(str(exc)) from exc
    directory = os.path.dirname(path)
    if directory:
        os.makedirs(directory, exist_ok=True)
    with open(path, "wb") as fh:
        fh.write(data)
    return path


# ── PDF ───────────────────────────────────────────────────────────────────
# A CJK-capable TrueType font (installed in the image) so non-Latin CVs render.
_CJK_FONT_NAME = "HookAI-CJK"
_CJK_FONT_PATHS = [
    "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
    "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
]
_cjk_state: "str | bool | None" = None  # None=unknown, str=font name, False=unavailable


# The non-Latin-1 code points WinAnsi/CP1252 (Times' default encoding) DOES cover —
# dashes, curly quotes, bullet, ellipsis, €, ™, etc. — so they must NOT trigger the
# CJK font (which would drop the serif look on ordinary resumes).
_WINANSI_EXTRA = {
    0x20AC, 0x201A, 0x0192, 0x201E, 0x2026, 0x2020, 0x2021, 0x02C6, 0x2030, 0x0160,
    0x2039, 0x0152, 0x017D, 0x2018, 0x2019, 0x201C, 0x201D, 0x2022, 0x2013, 0x2014,
    0x02DC, 0x2122, 0x0161, 0x203A, 0x0153, 0x017E, 0x0178,
}


def _needs_unicode_font(text: str) -> bool:
    """True if any character can't be drawn by Times/WinAnsi (CJK, Cyrillic,
    Vietnamese, …) and needs an embedded Unicode font to avoid glyph boxes."""
    for ch in text:
        o = ord(ch)
        if o in (0x09, 0x0A, 0x0D):  # tab / newline
            continue
        if 0x20 <= o <= 0x7E or 0xA0 <= o <= 0xFF or o in _WINANSI_EXTRA:
            continue
        return True
    return False


def _cjk_font() -> "str | None":
    """Register (once) an embedded Unicode font; None if none is installed."""
    global _cjk_state
    if _cjk_state is not None:
        return _cjk_state or None
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        for path in _CJK_FONT_PATHS:
            if os.path.exists(path):
                pdfmetrics.registerFont(TTFont(_CJK_FONT_NAME, path, subfontIndex=0))
                # Map bold/italic to the same face so <b>/<i> markup doesn't error
                # (this font has one weight — hierarchy stays via size + section rules).
                pdfmetrics.registerFontFamily(
                    _CJK_FONT_NAME, normal=_CJK_FONT_NAME, bold=_CJK_FONT_NAME,
                    italic=_CJK_FONT_NAME, boldItalic=_CJK_FONT_NAME)
                _cjk_state = _CJK_FONT_NAME
                return _CJK_FONT_NAME
    except Exception:
        pass
    _cjk_state = False
    return None


def _pdf_escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _pdf_inline(text: str) -> str:
    return _BOLD_RE.sub(r"<b>\1</b>", _pdf_escape(text))


def render_pdf_bytes(markdown: str) -> bytes:
    """Black-and-white serif resume as an in-memory .pdf (reportlab).

    Real text-searchable PDF (not a screenshot), so it stays ATS-parseable."""
    try:
        from reportlab.lib.enums import TA_CENTER, TA_RIGHT
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            HRFlowable, ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
        )
    except ImportError as exc:
        raise DocgenError("PDF export needs reportlab (pip install reportlab)") from exc

    serif, serif_b, serif_i = "Times-Roman", "Times-Bold", "Times-Italic"
    if _needs_unicode_font(markdown):
        cjk = _cjk_font()
        if cjk:  # one face covers CJK + Latin; keep hierarchy via size + section rules
            serif = serif_b = serif_i = cjk
    name = ParagraphStyle("Name", fontName=serif, fontSize=22, alignment=TA_CENTER, spaceAfter=2, leading=25)
    contact = ParagraphStyle("Contact", fontName=serif, fontSize=9.5, alignment=TA_CENTER, spaceAfter=6, leading=12)
    section = ParagraphStyle("Section", fontName=serif_b, fontSize=12, spaceBefore=9, spaceAfter=1, leading=14)
    body = ParagraphStyle("Body", fontName=serif, fontSize=10.5, leading=13, spaceAfter=2)
    left_bold = ParagraphStyle("LBold", fontName=serif_b, fontSize=10.5, leading=13)
    left_ital = ParagraphStyle("LItal", fontName=serif_i, fontSize=10.5, leading=13)
    right_reg = ParagraphStyle("RReg", fontName=serif, fontSize=10.5, leading=13, alignment=TA_RIGHT)
    right_ital = ParagraphStyle("RItal", fontName=serif_i, fontSize=10.5, leading=13, alignment=TA_RIGHT)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=LETTER, title="ATS CV",
        topMargin=0.5 * inch, bottomMargin=0.5 * inch, leftMargin=0.7 * inch, rightMargin=0.7 * inch,
    )
    avail = doc.width
    story: list = []
    bullets: list[str] = []
    contact_parts: list[str] = []
    seen_section = False

    def flush_bullets() -> None:
        if bullets:
            story.append(ListFlowable(
                [ListItem(Paragraph(_pdf_inline(b), body), leftIndent=10) for b in bullets],
                bulletType="bullet", leftIndent=13, bulletFontSize=7, start="•",
            ))
            story.append(Spacer(1, 3))
            bullets.clear()

    def flush_contact() -> None:
        if contact_parts:
            story.append(Paragraph(" | ".join(_pdf_escape(c) for c in contact_parts), contact))
            contact_parts.clear()

    def two_col(left_text: str, right_text: str, italic: bool) -> None:
        left_style = left_ital if italic else left_bold
        right_style = right_ital if italic else right_reg
        row = Table(
            [[Paragraph(_pdf_escape(left_text), left_style), Paragraph(_pdf_escape(right_text), right_style)]],
            colWidths=[avail * 0.68, avail * 0.32],
        )
        row.setStyle(TableStyle([
            ("LEFTPADDING", (0, 0), (-1, -1), 0), ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0), ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(row)

    for block in parse_cv_markdown(markdown):
        if block.kind == "bullet":
            bullets.append(block.text)
            continue
        flush_bullets()

        if block.kind == "h1":
            story.append(Paragraph(_pdf_escape(_strip_bold(block.text)), name))
        elif block.kind == "para" and not seen_section:
            contact_parts.append(block.text)
        elif block.kind == "h2":
            flush_contact()
            seen_section = True
            story.append(Paragraph(_strip_bold(block.text).upper(), section))
            story.append(HRFlowable(width="100%", thickness=0.75, color="black", spaceBefore=1, spaceAfter=5))
        elif block.kind == "entry_head":
            two_col(block.text, block.right, italic=False)
        elif block.kind == "entry_sub":
            two_col(block.text, block.right, italic=True)
        elif block.kind == "h3":
            story.append(Paragraph("<b>" + _pdf_escape(block.text) + "</b>", body))
        else:  # para (summary / skills)
            story.append(Paragraph(_pdf_inline(block.text), body))
    flush_bullets()
    flush_contact()

    doc.build(story)
    return buf.getvalue()


def write_pdf_from_markdown(path: str, markdown: str) -> str:
    try:
        data = render_pdf_bytes(markdown)
    except DocgenError as exc:
        raise SystemExit(str(exc)) from exc
    directory = os.path.dirname(path)
    if directory:
        os.makedirs(directory, exist_ok=True)
    with open(path, "wb") as fh:
        fh.write(data)
    return path
